# Note: You need python-docx library to make this script work.
import os, re, docx
import numpy as np
from docx.oxml.ns import qn
from docx.shared import Cm

# List all files ending with .xyz in the path of the script.
fullXYZPaths, fullLOGPaths, logNames = [], [], []
for root, dirs, files in os.walk("./"):
    for file in files:
        path = os.path.join(root, file)
        if re.search('M062X_def2tzvpp.+\.xyz', path): fullXYZPaths.append(path)
        if re.search('.+Pathway.+?\.log', path):
            fullLOGPaths.append(path)
            logNames.append(re.sub('.+/(.+)\.log', r'\1', fullLOGPaths[-1]))

# Create a Word document
doc = docx.Document()
section = doc.sections[0]
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'2')
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

for i in range(len(fullXYZPaths)):
    # Read xyz file
    # Format:
    # Line 1: [Number of atoms]
    # Line 2: [Molecule name]
    # Line 3-: [Atom type] [Pos x] [Pos y] [Pos z]
    with open(fullXYZPaths[i], 'r') as file: xyzFileLines = file.readlines()

    # Read log file
    logFileLines = []
    freq = ""
    # print(fullXYZPaths[i])
    # input()
    for j in range(len(fullLOGPaths)):
        if re.search(f'/{logNames[j]}-', fullXYZPaths[i]):
            # print(fullXYZPaths[i])
            # print(fullLOGPaths[j])
            # print(logNames[j])
            # input()
            with open(fullLOGPaths[j], 'r') as file: logFileLines = file.readlines()
            break
    # Search for imaginary frequency in log file.
    # If frequency not found, freq will be a string, else it will be a floating number.
    for line in logFileLines:
        freq = re.sub('.*Frequencies.+?(-\d+\.\d+).+', r'\1', line)
        if re.match('-\d+\.\d+', freq):
            freq = float(freq)
            break

    # Remove ".pdb" to the name.
    xyzFileLines[1] = re.sub('(.+?)(\..+)', r'\1', xyzFileLines[1])

    # Convert fileLines to text.
    text2 = ''.join(xyzFileLines[2:])

    # Add a paragraph to the Word document
    paragraph = doc.add_paragraph()

    # Add some formatting to the paragraph
    paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.space_after = 0

    # Add a run to the paragraph
    paragraph.add_run(xyzFileLines[1]).bold = True
    # If xyz file is a transition state (TS), add imaginary frequency.
    if isinstance(freq, float):
        paragraph.add_run(f'Imaginary frequency: {freq:.4f} cm').font.size = docx.shared.Pt(12)
        paragraph.add_run('-1').font.superscript = True
        paragraph.add_run('\n')
    run = paragraph.add_run(text2)

    # Add some formatting to the run
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)

    # # Add another paragraph (left blank for an empty line)
    # doc.add_paragraph()

# Save the document
doc.save("positions.docx")

